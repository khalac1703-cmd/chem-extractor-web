import io, os, json, html, tempfile
import streamlit as st
from typing import List
from bs4 import BeautifulSoup
from chem_rules import chem_transform_v3, clean_footer
# --- Option sanitizer (ABCD) ---
import re
from typing import List

# A./B./C./D./Đ. (chấp nhận có/không khoảng trắng; dấu . hoặc ))
OPT_LABEL = r'(?:[A-DĐ])\s*[.)]'

# Tách nhiều nhãn A/B/C/D nếu chúng xuất hiện nối tiếp nhau trên cùng 1 dòng
# Ví dụ: "A. 10   B. 20   C. 30" -> tách thành 3 dòng
MULTI_OPT_SPLIT_RE = re.compile(r'\s+(?=([A-DĐ])\s*[.)]\s+)')

# Nhận diện 1 dòng đáp án bắt đầu bằng A./B./C./D./Đ.
LINE_OPT_RE = re.compile(r'^\s*([A-DĐ])\s*[.)]\s*(.*)$')

def sanitize_lines_for_options(lines: List[str]) -> List[str]:
    out: List[str] = []

    # 1) Tách các nhãn nếu có nhiều đáp án trên cùng một dòng
    expanded: List[str] = []
    for ln in lines:
        ln = ln.rstrip().rstrip("\u00A0")  # bỏ space/NBSP cuối dòng
        # Nếu có từ 2 nhãn trở lên trên 1 dòng
        if re.search(OPT_LABEL + r'.*\s+' + OPT_LABEL, ln):
            parts = MULTI_OPT_SPLIT_RE.split(ln)
            buf = ""
            for p in parts:
                if re.fullmatch(r'[A-DĐ]', p or ""):
                    if buf.strip():
                        expanded.append(buf.strip())
                    buf = p + ". "
                else:
                    buf += (p or "")
            if buf.strip():
                expanded.append(buf.strip())
        else:
            expanded.append(ln)

    # 2) Ép định dạng chuẩn morat: "A. {body}" (đúng 1 khoảng trắng sau "A.")
    tmp: List[str] = []
    for ln in expanded:
        m = LINE_OPT_RE.match(ln)
        if m:
            label, body = m.group(1), m.group(2).lstrip()
            ln = f"{label}. {body}"
        tmp.append(ln)

    # 3) Xoá dòng trống giữa các đáp án liên tiếp
    def is_option(s: str) -> bool:
        return bool(LINE_OPT_RE.match(s))

    for i, ln in enumerate(tmp):
        if ln.strip() == "":
            prev_is_opt = i > 0 and is_option(tmp[i-1])
            next_is_opt = i+1 < len(tmp) and is_option(tmp[i+1])
            if prev_is_opt and next_is_opt:
                continue  # bỏ dòng trống xen giữa các đáp án
        out.append(ln)

    return out

# -------- PDF text extraction --------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    from pdfminer.high_level import extract_text
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        return extract_text(tmp_path)
    finally:
        os.unlink(tmp_path)

def to_html(lines: List[str], normalize_arrows: bool) -> str:
    parts = ['<html><head><meta charset="utf-8"><title>Chem</title></head>'
             '<body style="font-family:Times New Roman,serif; line-height:1.35; font-size:14px; white-space:pre-wrap;">']
    for ln in lines:
        ln = ln.rstrip()   # <<< quan trọng: bỏ space ở cuối dòng
        parts.append(f"<div>{chem_transform_v3(ln, normalize_arrows)}</div>")
    parts.append("</body></html>")
    return "\n".join(parts)
# ---- HTML -> DOCX (giữ sub/sup + basic bold/italic nếu có trong HTML) ----
def html_to_docx(chem_html: str) -> bytes:
    from docx import Document
    doc = Document()
    soup = BeautifulSoup(chem_html, "lxml")

    def add_nodes(parent, paragraph):
        for node in parent.children:
            if isinstance(node, str):
                txt = node.rstrip()          # <<< bỏ space cuối
                if txt: paragraph.add_run(txt)
            else:
                tag = node.name.lower()
                if tag == "sub":
                    r = paragraph.add_run(node.get_text().rstrip()); r.font.subscript = True
                elif tag == "sup":
                    r = paragraph.add_run(node.get_text().rstrip()); r.font.superscript = True
                elif tag in ("b","strong"):
                    r = paragraph.add_run(node.get_text().rstrip()); r.bold = True
                elif tag in ("i","em"):
                    r = paragraph.add_run(node.get_text().rstrip()); r.italic = True
                else:
                    if node.contents:
                        add_nodes(node, paragraph)
                    else:
                        t = node.get_text().rstrip()
                        if t: paragraph.add_run(t)

    for div in soup.select("body > div, p"):
        p = doc.add_paragraph()
        add_nodes(div, p)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ---------------- UI ----------------
st.set_page_config(page_title="Chem Extractor Web", page_icon="🧪", layout="wide")
st.title("🧪 Chem Extractor Web — PDF → HTML/DOCX")
st.caption("Giữ chỉ số dưới/trên (chem-aware), công tắc chuẩn hoá mũi tên, bỏ footer/số trang, xuất Word (.docx)")

colA, colB = st.columns(2)
normalize = colA.toggle("Chuẩn hoá mũi tên ASCII → Unicode", value=True)
drop_footer = colB.toggle("Bỏ footer / số trang", value=True)

uploaded = st.file_uploader("Tải đề (.pdf — ưu tiên PDF text-based; nếu .docx hãy lưu ra PDF rồi tải)", type=["pdf","docx"])
if uploaded is not None:
    if uploaded.type.endswith("pdf"):
        text = extract_text_from_pdf(uploaded.getvalue())
    else:
        st.warning("Bản web hiện hỗ trợ PDF tốt nhất. Hãy lưu Word ra PDF rồi tải lên.")
        text = ""

    if not text.strip():
        st.error("Không trích xuất được text. Kiểm tra file có phải scan/ảnh không.")
    else:
        lines = text.replace("\r\n","\n").replace("\r","\n").split("\n")
        if drop_footer:
            lines, removed = clean_footer(lines)
            st.write("Removed lines:", removed)
         # Hợp nhất các đoạn văn bị ngắt dòng giữa chừng (Nomex, mô tả dài,...)
         from chem_rules import merge_broken_paragraphs
         lines = merge_broken_paragraphs(lines)
         lines = sanitize_lines_for_options(lines)
            with st.expander("Dòng đã loại bỏ (footer/số trang)"):
                st.json(removed)

        html_out = to_html(lines, normalize_arrows=normalize)

        st.subheader("Xem trước HTML (chem-aware)")
        st.components.v1.html(html_out, height=600, scrolling=True)

        st.download_button("⬇️ Tải HTML", data=html_out.encode("utf-8"),
                           file_name="chem_exam.html", mime="text/html")

        docx_bytes = html_to_docx(html_out)
        st.download_button("⬇️ Tải Word (.docx)", data=docx_bytes,
                           file_name="chem_exam.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
